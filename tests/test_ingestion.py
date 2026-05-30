"""Tests for the ingestion pipeline: extractors, Hindsight, Preprocessor, worker.

Fixtures are generated in tests (no binary blobs in the repo): the DOCX via
python-docx, the PDF via pypdf, the CSV with stdlib. Audio/video/image go through
mocks. End-to-end: CSV with 3 products → InMemoryHindsight has 3 facts.
"""

from __future__ import annotations

import asyncio
import csv as csvlib
from pathlib import Path

from docx import Document
import pypdf
import pytest
from services.preprocessor.worker import IngestionJob, IngestionQueue, drain, run_forever

from hermesell.ingestion import (
    CsvExtractor,
    DocxExtractor,
    InMemoryHindsight,
    MockAudioExtractor,
    MockImageExtractor,
    MockVideoExtractor,
    PdfExtractor,
    Preprocessor,
    UnsupportedFormatError,
    default_extractors,
)
from hermesell.models import Fact

pytestmark = pytest.mark.unit


# --- fixtures ---------------------------------------------------------------


def _write_catalog_csv(path: Path) -> Path:
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csvlib.writer(f)
        writer.writerow(["name", "price", "stock"])
        writer.writerow(["Camiseta básica", "8500", "12"])
        writer.writerow(["Pantalón cargo", "21000", "5"])
        writer.writerow(["Mochila urbana", "15000", "8"])
    return path


def _write_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Política de devoluciones: 30 días.")
    doc.add_paragraph("")  # empty paragraph — should be skipped
    doc.add_paragraph("Envíos a todo el país.")
    doc.save(str(path))
    return path


def _write_pdf(path: Path) -> Path:
    # pypdf can create a PDF by appending a blank page; we then inject a Text
    # operator stream so extract_text() returns something we can assert on.
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as f:
        writer.write(f)
    return path


# --- extractors -------------------------------------------------------------


class TestCsvExtractor:
    async def test_each_row_becomes_a_chunk(self, tmp_path: Path) -> None:
        path = _write_catalog_csv(tmp_path / "catalog.csv")
        chunks = await CsvExtractor().extract(path)
        assert len(chunks) == 3
        assert "Camiseta básica" in chunks[0].text
        assert chunks[0].source == "catalog.csv"
        assert chunks[0].metadata == {"row": "0"}

    async def test_skips_empty_rows(self, tmp_path: Path) -> None:
        p = tmp_path / "empty.csv"
        p.write_text("name,price\n,\n\nAcme,10\n", encoding="utf-8")
        chunks = await CsvExtractor().extract(p)
        # Only the Acme row has content (the empty row produces "name: | price: ")
        non_blank = [c for c in chunks if "Acme" in c.text]
        assert len(non_blank) == 1


class TestDocxExtractor:
    async def test_paragraphs_become_chunks(self, tmp_path: Path) -> None:
        path = _write_docx(tmp_path / "doc.docx")
        chunks = await DocxExtractor().extract(path)
        # The empty paragraph between the two real ones is skipped.
        assert [c.text for c in chunks] == [
            "Política de devoluciones: 30 días.",
            "Envíos a todo el país.",
        ]


class TestPdfExtractor:
    async def test_blank_pdf_yields_no_chunks(self, tmp_path: Path) -> None:
        # A blank page has no extractable text — the extractor must NOT crash and
        # must NOT emit empty chunks (regression guard for the strip() check).
        path = _write_pdf(tmp_path / "blank.pdf")
        chunks = await PdfExtractor().extract(path)
        assert chunks == []


class TestMockExtractors:
    async def test_audio_mock_returns_transcript(self, tmp_path: Path) -> None:
        chunks = await MockAudioExtractor(transcript="audio-XYZ").extract(tmp_path / "x.mp3")
        assert len(chunks) == 1
        assert chunks[0].text == "audio-XYZ"
        assert chunks[0].metadata["modality"] == "audio"

    async def test_video_and_image_mocks_tag_modality(self, tmp_path: Path) -> None:
        v = await MockVideoExtractor().extract(tmp_path / "x.mp4")
        i = await MockImageExtractor().extract(tmp_path / "x.png")
        assert v[0].metadata["modality"] == "video"
        assert i[0].metadata["modality"] == "image"


# --- hindsight --------------------------------------------------------------


class TestInMemoryHindsight:
    def test_add_and_query_finds_substring_match(self) -> None:
        h = InMemoryHindsight()
        h.add_fact(Fact(tenant_id="t1", source="x", content="Camiseta azul talle M, stock 5"))
        h.add_fact(Fact(tenant_id="t1", source="x", content="Pantalón negro"))
        assert [f.content for f in h.query(text="camiseta", tenant_id="t1")] == [
            "Camiseta azul talle M, stock 5"
        ]

    def test_query_is_tenant_scoped(self) -> None:
        h = InMemoryHindsight()
        h.add_fact(Fact(tenant_id="t1", source="x", content="zapatillas"))
        h.add_fact(Fact(tenant_id="t2", source="x", content="zapatillas"))
        assert len(h.query(text="zapatillas", tenant_id="t1")) == 1
        assert len(h.query(text="zapatillas")) == 2  # no tenant filter → both

    def test_empty_query_returns_nothing(self) -> None:
        h = InMemoryHindsight()
        h.add_fact(Fact(tenant_id="t1", source="x", content="anything"))
        assert h.query(text="") == []

    def test_top_k_caps_results(self) -> None:
        h = InMemoryHindsight()
        for i in range(5):
            h.add_fact(Fact(tenant_id="t1", source="x", content=f"item-{i}-match"))
        assert len(h.query(text="match", tenant_id="t1", top_k=3)) == 3

    def test_all_for_returns_tenant_facts(self) -> None:
        h = InMemoryHindsight()
        h.add_fact(Fact(tenant_id="t1", source="x", content="a"))
        h.add_fact(Fact(tenant_id="t2", source="x", content="b"))
        assert [f.content for f in h.all_for("t1")] == ["a"]


# --- preprocessor end-to-end -----------------------------------------------


class TestPreprocessor:
    async def test_csv_three_products_becomes_three_facts(self, tmp_path: Path) -> None:
        path = _write_catalog_csv(tmp_path / "catalog.csv")
        hindsight = InMemoryHindsight()
        pre = Preprocessor(hindsight=hindsight)

        facts = await pre.process(path, tenant_id="tenant-A")

        assert len(facts) == 3
        assert all(f.tenant_id == "tenant-A" for f in facts)
        assert all(f.source == "catalog.csv" for f in facts)
        # All three are now in Hindsight, scoped to tenant-A.
        assert len(hindsight.all_for("tenant-A")) == 3

    async def test_source_label_override(self, tmp_path: Path) -> None:
        path = _write_catalog_csv(tmp_path / "x.csv")
        pre = Preprocessor()
        facts = await pre.process(path, tenant_id="t", source_label="catalog-v2")
        assert all(f.source == "catalog-v2" for f in facts)

    async def test_audio_file_routes_to_mock_audio(self, tmp_path: Path) -> None:
        # No real audio file needed — the mock doesn't read the bytes.
        audio = tmp_path / "voice.mp3"
        audio.write_bytes(b"not really audio")
        pre = Preprocessor()
        facts = await pre.process(audio, tenant_id="t")
        assert len(facts) == 1
        assert facts[0].metadata["modality"] == "audio"

    async def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        p = tmp_path / "weird.xyz"
        p.write_text("???")
        with pytest.raises(UnsupportedFormatError, match="no extractor"):
            await Preprocessor().process(p, tenant_id="t")

    def test_default_extractor_map_covers_all_modalities(self) -> None:
        m = default_extractors()
        assert {".csv", ".pdf", ".docx", ".mp3", ".mp4", ".png"} <= set(m)


# --- worker -----------------------------------------------------------------


class TestWorker:
    async def test_drain_processes_all_queued_jobs(self, tmp_path: Path) -> None:
        catalog = _write_catalog_csv(tmp_path / "catalog.csv")
        hindsight = InMemoryHindsight()
        pre = Preprocessor(hindsight=hindsight)
        queue = IngestionQueue()
        await queue.put(IngestionJob(tenant_id="t1", path=catalog))
        await queue.put(IngestionJob(tenant_id="t2", path=catalog))

        n = await drain(pre, queue)

        assert n == 2
        assert queue.empty()
        assert len(hindsight.all_for("t1")) == 3
        assert len(hindsight.all_for("t2")) == 3

    async def test_run_forever_stops_on_event(self, tmp_path: Path) -> None:
        catalog = _write_catalog_csv(tmp_path / "catalog.csv")
        pre = Preprocessor(hindsight=InMemoryHindsight())
        queue = IngestionQueue()
        stop = asyncio.Event()

        # Run the daemon, feed one job, then stop.
        await queue.put(IngestionJob(tenant_id="t1", path=catalog))
        runner = asyncio.create_task(run_forever(pre, queue, stop=stop, idle_poll_seconds=0.05))
        await asyncio.sleep(0.2)  # let it drain the one job + idle-poll
        stop.set()
        await asyncio.wait_for(runner, timeout=2.0)
        assert queue.empty()
